"""
backup_utils.py — سیستم پشتیبان‌گیری کامل
شامل: تولید Excel، تولید Word، بکاپ خودکار با تایمر
"""
import io
import os
import logging
from datetime import datetime, timedelta

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document

import database as db
from helpers import now_shamsi, today_shamsi

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────
# جمع‌آوری داده برای بکاپ
# ──────────────────────────────────────────────────────────────
async def gather_data(period: str = "all") -> dict:
    from datetime import timedelta
    now = datetime.now()

    matches = await db.get_matches_by_filter(period)
    players = await db.get_all_players()
    tournaments = await db.get_all_tournaments()
    classes = await db.get_all_classes()
    admins = await db.get_all_admins()

    return {
        "players": players,
        "matches": matches,
        "tournaments": tournaments,
        "classes": classes,
        "admins": admins,
        "generated_at": now_shamsi(),
        "period": period,
        "total_players": len(players),
        "total_matches": len(matches),
    }


# ──────────────────────────────────────────────────────────────
# ابزارهای استایل Excel
# ──────────────────────────────────────────────────────────────
def _header_style():
    return {
        "font": Font(bold=True, color="FFFFFF", size=11, name="Arial"),
        "fill": PatternFill("solid", fgColor="1F3864"),
        "alignment": Alignment(horizontal="center", vertical="center", wrap_text=True),
        "border": Border(
            left=Side(style="thin", color="AAAAAA"),
            right=Side(style="thin", color="AAAAAA"),
            top=Side(style="thin", color="AAAAAA"),
            bottom=Side(style="thin", color="AAAAAA"),
        )
    }

def _row_style(even: bool):
    return {
        "fill": PatternFill("solid", fgColor="EBF1FF" if even else "FFFFFF"),
        "alignment": Alignment(horizontal="right", vertical="center", wrap_text=True),
        "border": Border(
            left=Side(style="thin", color="DDDDDD"),
            right=Side(style="thin", color="DDDDDD"),
            top=Side(style="thin", color="DDDDDD"),
            bottom=Side(style="thin", color="DDDDDD"),
        )
    }

def _apply_style(cell, style: dict):
    for attr, val in style.items():
        setattr(cell, attr, val)

def _write_header(ws, headers: list):
    style = _header_style()
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        _apply_style(cell, style)
    ws.row_dimensions[1].height = 22

def _write_row(ws, row_num: int, values: list):
    style = _row_style(row_num % 2 == 0)
    for col, v in enumerate(values, 1):
        cell = ws.cell(row=row_num, column=col, value=str(v) if v is not None else "")
        _apply_style(cell, style)

def _auto_width(ws, min_width=12, max_width=35):
    for col in ws.columns:
        max_len = min_width
        col_letter = col[0].column_letter
        for cell in col:
            try:
                val = str(cell.value or "")
                max_len = min(max(max_len, len(val) + 2), max_width)
            except Exception:
                pass
        ws.column_dimensions[col_letter].width = max_len


# ──────────────────────────────────────────────────────────────
# تولید فایل Excel
# ──────────────────────────────────────────────────────────────
async def generate_excel_backup(period: str = "all") -> io.BytesIO:
    data = await gather_data(period)
    wb = openpyxl.Workbook()

    # ─── شیت اطلاعات ─────────────────────────────────────────
    ws_info = wb.active
    ws_info.title = "اطلاعات بکاپ"
    ws_info.sheet_view.rightToLeft = True
    info_rows = [
        ("عنوان", "سیستم فرماندهی شطرنج — گزارش پشتیبان"),
        ("تاریخ تهیه", data["generated_at"]),
        ("بازه زمانی", {
            "today": "امروز", "week": "این هفته",
            "month": "این ماه", "all": "از ابتدا"
        }.get(period, period)),
        ("تعداد بازیکنان", data["total_players"]),
        ("تعداد مسابقات", data["total_matches"]),
        ("تعداد تورنمنت‌ها", len(data["tournaments"])),
        ("تعداد کلاس‌ها", len(data["classes"])),
    ]
    for i, (k, v) in enumerate(info_rows, 1):
        ws_info.cell(row=i, column=1, value=k).font = Font(bold=True, name="Arial")
        ws_info.cell(row=i, column=2, value=str(v)).font = Font(name="Arial")
    ws_info.column_dimensions["A"].width = 22
    ws_info.column_dimensions["B"].width = 38

    # ─── شیت بازیکنان ────────────────────────────────────────
    ws_p = wb.create_sheet("بازیکنان")
    ws_p.sheet_view.rightToLeft = True
    headers_p = ["ردیف", "نام کامل", "کلاس", "وضعیت", "برد", "مساوی", "باخت",
                 "اخطار", "برتر", "ویژه", "تاریخ ثبت"]
    _write_header(ws_p, headers_p)
    status_map = {
        "active": "فعال", "suspended": "تعلیق",
        "kicked": "اخراج", "eliminated": "حذف"
    }
    for i, p in enumerate(data["players"], 1):
        _write_row(ws_p, i + 1, [
            i,
            p["full_name"] or "",
            p["class_name"] or "" if p["class_name"] else "",
            status_map.get(p["status"], p["status"]),
            p["wins"], p["draws"], p["losses"],
            p["warnings"],
            "بله" if p["is_elite"] else "خیر",
            "بله" if p["is_special"] else "خیر",
            str(p["created_at"] or "")[:10],
        ])
    _auto_width(ws_p)

    # ─── شیت مسابقات ─────────────────────────────────────────
    ws_m = wb.create_sheet("مسابقات")
    ws_m.sheet_view.rightToLeft = True
    headers_m = ["ردیف", "بازیکن سفید", "بازیکن سیاه", "نتیجه",
                 "علت تساوی", "تاریخ مسابقه", "ثبت‌کننده", "زمان ثبت"]
    _write_header(ws_m, headers_m)
    result_fa = {"white": "برد سفید", "black": "برد سیاه",
                 "draw": "تساوی", None: "بدون نتیجه", "": "بدون نتیجه"}
    for i, m in enumerate(data["matches"], 1):
        _write_row(ws_m, i + 1, [
            i,
            m["white_name"] or "",
            m["black_name"] or "",
            result_fa.get(m["result"], str(m["result"] or "")),
            m["draw_reason"] or "",
            str(m["match_date"] or ""),
            str(m["created_by"] or ""),
            str(m["created_at"] or "")[:16],
        ])
    _auto_width(ws_m)

    # ─── شیت تورنمنت‌ها ──────────────────────────────────────
    ws_t = wb.create_sheet("تورنمنت‌ها")
    ws_t.sheet_view.rightToLeft = True
    headers_t = ["ردیف", "نام", "وضعیت", "پیش‌فرض", "تاریخ ایجاد"]
    _write_header(ws_t, headers_t)
    for i, t in enumerate(data["tournaments"], 1):
        _write_row(ws_t, i + 1, [
            i, t["name"], t["status"],
            "بله" if t["is_default"] else "خیر",
            str(t["created_at"] or "")[:10],
        ])
    _auto_width(ws_t)

    # ─── شیت کلاس‌ها ─────────────────────────────────────────
    ws_c = wb.create_sheet("کلاس‌ها")
    ws_c.sheet_view.rightToLeft = True
    headers_c = ["ردیف", "نام کلاس", "تاریخ ثبت"]
    _write_header(ws_c, headers_c)
    for i, c in enumerate(data["classes"], 1):
        _write_row(ws_c, i + 1, [i, c["name"], str(c["created_at"] or "")[:10]])
    _auto_width(ws_c)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────────────────────
# تولید فایل Word
# ──────────────────────────────────────────────────────────────
async def generate_word_backup(period: str = "all") -> io.BytesIO:
    data = await gather_data(period)
    doc = Document()

    # تنظیم فونت پیش‌فرض
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # عنوان
    title = doc.add_heading("سیستم فرماندهی شطرنج — گزارش پشتیبان", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"📅 تاریخ تهیه: {data['generated_at']}")
    doc.add_paragraph(f"📊 بازه: {period} | بازیکنان: {data['total_players']} | مسابقات: {data['total_matches']}")

    p = doc.add_paragraph()
    p.add_run("─" * 50)

    # ─── بازیکنان ────────────────────────────────────────────
    doc.add_heading("👤 بازیکنان", level=1)
    if data["players"]:
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["نام", "کلاس", "وضعیت", "برد", "مساوی", "باخت", "اخطار", "تاریخ"]):
            hdr[i].text = h
            try:
                hdr[i].paragraphs[0].runs[0].font.bold = True
            except IndexError:
                run = hdr[i].paragraphs[0].add_run(h)
                run.font.bold = True
                hdr[i].text = ""

        status_map = {"active": "فعال", "suspended": "تعلیق",
                      "kicked": "اخراج", "eliminated": "حذف"}
        for p_row in data["players"]:
            row = table.add_row().cells
            row[0].text = p_row["full_name"] or ""
            row[1].text = (p_row["class_name"] or "") if p_row["class_name"] else ""
            row[2].text = status_map.get(p_row["status"], p_row["status"])
            row[3].text = str(p_row["wins"])
            row[4].text = str(p_row["draws"])
            row[5].text = str(p_row["losses"])
            row[6].text = str(p_row["warnings"])
            row[7].text = str(p_row["created_at"] or "")[:10]
    else:
        doc.add_paragraph("هیچ بازیکنی ثبت نشده.")

    doc.add_paragraph()

    # ─── مسابقات ─────────────────────────────────────────────
    doc.add_heading("♟️ مسابقات", level=1)
    if data["matches"]:
        table2 = doc.add_table(rows=1, cols=5)
        table2.style = "Table Grid"
        hdr2 = table2.rows[0].cells
        for i, h in enumerate(["سفید", "سیاه", "نتیجه", "علت تساوی", "تاریخ"]):
            hdr2[i].text = h
            try:
                hdr2[i].paragraphs[0].runs[0].font.bold = True
            except IndexError:
                pass

        result_fa = {"white": "برد سفید", "black": "برد سیاه",
                     "draw": "تساوی", None: "—"}
        for m in data["matches"]:
            row = table2.add_row().cells
            row[0].text = m["white_name"] or ""
            row[1].text = m["black_name"] or ""
            row[2].text = result_fa.get(m["result"], str(m["result"] or ""))
            row[3].text = m["draw_reason"] or ""
            row[4].text = str(m["match_date"] or "")
    else:
        doc.add_paragraph("هیچ مسابقه‌ای ثبت نشده.")

    doc.add_paragraph()

    # ─── تورنمنت‌ها ───────────────────────────────────────────
    doc.add_heading("🏅 تورنمنت‌ها", level=1)
    for t in data["tournaments"]:
        doc.add_paragraph(
            f"• {t['name']} | {t['status']} | {'پیش‌فرض ✅' if t['is_default'] else ''}"
        )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────────────────────
# نام فایل استاندارد
# ──────────────────────────────────────────────────────────────
def backup_filename(fmt: str, period: str) -> str:
    ts = today_shamsi().replace("/", "-")
    period_fa = {
        "today": "امروز", "week": "هفته", "month": "ماه", "all": "کامل"
    }.get(period, period)
    ext = "xlsx" if fmt == "excel" else "docx"
    return f"chess_backup_{period_fa}_{ts}.{ext}"


# ──────────────────────────────────────────────────────────────
# ارسال بکاپ
# ──────────────────────────────────────────────────────────────
async def send_backup(bot, chat_id: int, period: str, fmt: str):
    """تهیه و ارسال فایل بکاپ"""
    try:
        if fmt == "excel":
            buf = await generate_excel_backup(period)
        else:
            buf = await generate_word_backup(period)

        filename = backup_filename(fmt, period)
        ts = now_shamsi()
        period_fa = {
            "today": "امروز", "week": "این هفته",
            "month": "این ماه", "all": "از ابتدا"
        }.get(period, period)

        caption = (
            f"💾 *بکاپ سیستم شطرنج*\n\n"
            f"📊 بازه: {period_fa}\n"
            f"📁 فرمت: {'Excel' if fmt == 'excel' else 'Word'}\n"
            f"⏱️ تهیه‌شده در: `{ts}`"
        )
        await bot.send_document(
            chat_id=chat_id,
            document=buf,
            filename=filename,
            caption=caption,
            parse_mode="Markdown"
        )
        logger.info(f"Backup sent to {chat_id}: {filename}")
        return True
    except Exception as e:
        logger.error(f"Backup failed: {e}")
        raise e


# ──────────────────────────────────────────────────────────────
# بکاپ خودکار (Job Queue)
# ──────────────────────────────────────────────────────────────
async def auto_backup_job(context):
    """این تابع توسط job queue صدا زده میشه"""
    from config import PISHVA_ID
    fmt = await db.get_setting("auto_backup_format", "excel")
    period = await db.get_setting("auto_backup_period", "all")
    logger.info(f"Auto backup running: {fmt} / {period}")
    try:
        await send_backup(context.bot, PISHVA_ID, period, fmt)
        await db.log_action(PISHVA_ID, "auto_backup", f"بکاپ خودکار: {fmt}/{period}")
    except Exception as e:
        logger.error(f"Auto backup job failed: {e}")
        try:
            await context.bot.send_message(
                chat_id=PISHVA_ID,
                text=f"❌ خطا در بکاپ خودکار:\n`{str(e)}`",
                parse_mode="Markdown"
            )
        except Exception:
            pass


def schedule_auto_backup(app, interval_hours: int = 24):
    """ثبت job بکاپ خودکار"""
    if not hasattr(app, 'job_queue') or app.job_queue is None:
        logger.warning("Job queue not available for auto backup")
        return False
    # حذف job قبلی اگر وجود داشت
    current = app.job_queue.get_jobs_by_name("auto_backup")
    for job in current:
        job.schedule_removal()
    # ثبت job جدید
    app.job_queue.run_repeating(
        auto_backup_job,
        interval=timedelta(hours=interval_hours),
        first=timedelta(seconds=30),
        name="auto_backup"
    )
    logger.info(f"Auto backup scheduled every {interval_hours} hours")
    return True
