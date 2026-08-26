"""
restore_utils.py — سیستم بازگردانی بکاپ (برعکسِ backup_utils.py)
شامل: تشخیص و خواندن فایل Excel/Word بکاپ، و وارد کردن محتوای آن به دیتابیس.

منطق کلی: هر فایلی که توسط backup_utils.py تولید شده (چه Excel چه Word)،
اینجا خونده و تحلیل می‌شه و دقیقاً همون داده‌ها (کلاس‌ها، بازیکنان،
تورنمنت‌ها، مسابقات) به سیستم بازگردانده می‌شن — یعنی مسیر معکوسِ بکاپ.
"""
import io
import logging

import openpyxl
from docx import Document

import database as db

logger = logging.getLogger(__name__)

STATUS_FA_TO_EN = {"فعال": "active", "تعلیق": "suspended", "اخراج": "kicked", "حذف": "eliminated"}
RESULT_FA_TO_EN = {"برد سفید": "white", "برد سیاه": "black", "تساوی": "draw"}
YESNO_FA_TO_BOOL = {"بله": 1, "خیر": 0}


def detect_format(filename: str) -> str:
    """فرمت فایل رو از روی پسوندش تشخیص می‌ده. خروجی: 'excel' یا 'word' یا None."""
    name = (filename or "").lower()
    if name.endswith(".xlsx") or name.endswith(".xls"):
        return "excel"
    if name.endswith(".docx"):
        return "word"
    return None


# ──────────────────────────────────────────────────────────────
# استخراج داده از Excel
# ──────────────────────────────────────────────────────────────
def parse_excel_backup(file_bytes: bytes) -> dict:
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    data = {"classes": [], "players": [], "tournaments": [], "matches": []}

    def sheet_rows(ws):
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        return [r for r in rows if any(c not in (None, "") for c in r)]

    if "کلاس‌ها" in wb.sheetnames:
        ws = wb["کلاس‌ها"]
        for row in sheet_rows(ws):
            # ["ردیف", "نام کلاس", "تاریخ ثبت"]
            name = str(row[1]).strip() if len(row) > 1 and row[1] else ""
            if name:
                data["classes"].append({"name": name})

    if "بازیکنان" in wb.sheetnames:
        ws = wb["بازیکنان"]
        for row in sheet_rows(ws):
            # ["ردیف","نام کامل","کلاس","وضعیت","برد","مساوی","باخت","اخطار","برتر","ویژه","تاریخ ثبت"]
            row = list(row) + [None] * (11 - len(row))
            full_name = str(row[1]).strip() if row[1] else ""
            if not full_name:
                continue
            data["players"].append({
                "full_name": full_name,
                "class_name": str(row[2]).strip() if row[2] else "",
                "status": STATUS_FA_TO_EN.get(str(row[3]).strip(), "active") if row[3] else "active",
                "wins": _to_int(row[4]),
                "draws": _to_int(row[5]),
                "losses": _to_int(row[6]),
                "warnings": _to_int(row[7]),
                "is_elite": YESNO_FA_TO_BOOL.get(str(row[8]).strip(), 0) if row[8] else 0,
                "is_special": YESNO_FA_TO_BOOL.get(str(row[9]).strip(), 0) if row[9] else 0,
                "created_at": str(row[10]).strip() if row[10] else None,
            })

    if "تورنمنت‌ها" in wb.sheetnames:
        ws = wb["تورنمنت‌ها"]
        for row in sheet_rows(ws):
            # ["ردیف", "نام", "وضعیت", "پیش‌فرض", "تاریخ ایجاد"]
            row = list(row) + [None] * (5 - len(row))
            name = str(row[1]).strip() if row[1] else ""
            if not name:
                continue
            data["tournaments"].append({
                "name": name,
                "status": str(row[2]).strip() if row[2] else "active",
                "is_default": YESNO_FA_TO_BOOL.get(str(row[3]).strip(), 0) if row[3] else 0,
            })

    if "مسابقات" in wb.sheetnames:
        ws = wb["مسابقات"]
        for row in sheet_rows(ws):
            # ["ردیف","بازیکن سفید","بازیکن سیاه","نتیجه","علت تساوی","تاریخ مسابقه","ثبت‌کننده","زمان ثبت"]
            row = list(row) + [None] * (8 - len(row))
            white = str(row[1]).strip() if row[1] else ""
            black = str(row[2]).strip() if row[2] else ""
            if not white and not black:
                continue
            data["matches"].append({
                "white_name": white,
                "black_name": black,
                "result": RESULT_FA_TO_EN.get(str(row[3]).strip()) if row[3] else None,
                "draw_reason": str(row[4]).strip() if row[4] else "",
                "match_date": str(row[5]).strip() if row[5] else "",
                "created_by_raw": str(row[6]).strip() if row[6] else "",
                "created_at": str(row[7]).strip() if row[7] else None,
            })

    return data


def _to_int(v) -> int:
    try:
        return int(v)
    except Exception:
        return 0


# ──────────────────────────────────────────────────────────────
# استخراج داده از Word
# ──────────────────────────────────────────────────────────────
def parse_word_backup(file_bytes: bytes) -> dict:
    doc = Document(io.BytesIO(file_bytes))
    data = {"classes": [], "players": [], "tournaments": [], "matches": []}

    # عناوین به همون ترتیبی که generate_word_backup می‌سازه ظاهر می‌شن:
    # 👤 بازیکنان (جدول ۱) ، ♟️ مسابقات (جدول ۲) ، 🏅 تورنمنت‌ها (پاراگراف)
    table_idx = 0
    for tbl in doc.tables:
        header_cells = [c.text.strip() for c in tbl.rows[0].cells]
        if not tbl.rows or len(tbl.rows) < 1:
            continue
        if "نام" in header_cells and "کلاس" in header_cells:
            for r in tbl.rows[1:]:
                vals = [c.text.strip() for c in r.cells]
                vals = vals + [""] * (8 - len(vals))
                full_name = vals[0]
                if not full_name:
                    continue
                data["players"].append({
                    "full_name": full_name,
                    "class_name": vals[1],
                    "status": STATUS_FA_TO_EN.get(vals[2], "active"),
                    "wins": _to_int(vals[3]),
                    "draws": _to_int(vals[4]),
                    "losses": _to_int(vals[5]),
                    "warnings": _to_int(vals[6]),
                    "is_elite": 0,
                    "is_special": 0,
                    "created_at": vals[7] or None,
                })
                if vals[1]:
                    data["classes"].append({"name": vals[1]})
        elif "سفید" in header_cells and "سیاه" in header_cells:
            for r in tbl.rows[1:]:
                vals = [c.text.strip() for c in r.cells]
                vals = vals + [""] * (5 - len(vals))
                white, black = vals[0], vals[1]
                if not white and not black:
                    continue
                data["matches"].append({
                    "white_name": white,
                    "black_name": black,
                    "result": RESULT_FA_TO_EN.get(vals[2]),
                    "draw_reason": vals[3],
                    "match_date": vals[4],
                    "created_by_raw": "",
                    "created_at": None,
                })
        table_idx += 1

    # تورنمنت‌ها به‌شکل پاراگراف‌های "• نام | وضعیت | پیش‌فرض ✅" ذخیره شدن
    in_tourn_section = False
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        if "تورنمنت" in txt and p.style and p.style.name and "Heading" in p.style.name:
            in_tourn_section = True
            continue
        if p.style and p.style.name and "Heading" in p.style.name:
            in_tourn_section = False
            continue
        if in_tourn_section and txt.startswith("•"):
            parts = [x.strip() for x in txt.lstrip("•").split("|")]
            parts = parts + [""] * (3 - len(parts))
            name = parts[0]
            if name:
                data["tournaments"].append({
                    "name": name,
                    "status": parts[1] or "active",
                    "is_default": 1 if "✅" in parts[2] else 0,
                })

    # حذف کلاس‌های تکراری
    seen = set()
    uniq_classes = []
    for c in data["classes"]:
        if c["name"] not in seen:
            seen.add(c["name"])
            uniq_classes.append(c)
    data["classes"] = uniq_classes

    return data


# ──────────────────────────────────────────────────────────────
# اعمال داده‌ها روی دیتابیس
# ──────────────────────────────────────────────────────────────
async def apply_restore(data: dict, admin_id: int) -> dict:
    """داده‌های استخراج‌شده رو در دیتابیس درج/به‌روزرسانی می‌کنه.
    خروجی: شمارش موارد جدید/به‌روزشده/خطاها."""
    counts = {
        "classes_new": 0,
        "players_new": 0, "players_updated": 0,
        "tournaments_new": 0, "tournaments_updated": 0,
        "matches_new": 0, "matches_skipped": 0,
        "errors": [],
    }

    # ۱) کلاس‌ها
    class_ids = {}
    for c in data.get("classes", []):
        try:
            existing = await db.get_class_by_name(c["name"])
            cid = await db.get_or_create_class(c["name"])
            class_ids[c["name"]] = cid
            if not existing:
                counts["classes_new"] += 1
        except Exception as e:
            counts["errors"].append(f"کلاس «{c['name']}»: {e}")

    # ۲) بازیکنان
    player_ids = {}
    for p in data.get("players", []):
        try:
            cid = None
            if p.get("class_name"):
                cid = class_ids.get(p["class_name"])
                if cid is None:
                    cid = await db.get_or_create_class(p["class_name"])
                    class_ids[p["class_name"]] = cid
            pid, created = await db.restore_upsert_player(
                full_name=p["full_name"], class_id=cid, status=p.get("status"),
                warnings=p.get("warnings", 0), is_elite=p.get("is_elite", 0),
                is_special=p.get("is_special", 0), wins=p.get("wins", 0),
                losses=p.get("losses", 0), draws=p.get("draws", 0),
                created_at=p.get("created_at"),
            )
            if pid:
                player_ids[p["full_name"].strip().lower()] = pid
                counts["players_new" if created else "players_updated"] += 1
        except Exception as e:
            counts["errors"].append(f"بازیکن «{p['full_name']}»: {e}")

    # ۳) تورنمنت‌ها
    tournament_id_by_name = {}
    for t in data.get("tournaments", []):
        try:
            tid, created = await db.get_or_create_tournament(
                t["name"], status=t.get("status", "active"), is_default=bool(t.get("is_default"))
            )
            tournament_id_by_name[t["name"]] = tid
            counts["tournaments_new" if created else "tournaments_updated"] += 1
        except Exception as e:
            counts["errors"].append(f"تورنمنت «{t['name']}»: {e}")

    default_tid = None
    default_row = await db.get_default_tournament()
    if default_row:
        default_tid = default_row["id"]

    # ۴) مسابقات — بازیکنانی که فقط در مسابقات ظاهر شدن (و در شیت بازیکنان نبودن) هم ساخته می‌شن
    for m in data.get("matches", []):
        try:
            white_id = await _resolve_player(m.get("white_name"), player_ids)
            black_id = await _resolve_player(m.get("black_name"), player_ids)
            if not white_id or not black_id:
                counts["matches_skipped"] += 1
                continue
            tid = default_tid
            if tid is None and tournament_id_by_name:
                tid = next(iter(tournament_id_by_name.values()))
            await db.insert_match_raw(
                white_id=white_id, black_id=black_id, result=m.get("result"),
                draw_reason=m.get("draw_reason") or None, match_date=m.get("match_date") or "",
                tournament_id=tid, created_by=admin_id, created_at=m.get("created_at"),
            )
            counts["matches_new"] += 1
        except Exception as e:
            counts["errors"].append(f"مسابقه «{m.get('white_name')} - {m.get('black_name')}»: {e}")

    return counts


async def _resolve_player(name: str, cache: dict):
    if not name:
        return None
    key = name.strip().lower()
    if key in cache:
        return cache[key]
    existing = await db.get_player_by_name(name)
    if existing:
        cache[key] = existing["id"]
        return existing["id"]
    pid, _ = await db.restore_upsert_player(full_name=name)
    if pid:
        cache[key] = pid
    return pid


def build_summary_text(counts: dict) -> str:
    lines = [
        f"🏷️ کلاس‌های جدید: {counts['classes_new']}",
        f"👤 بازیکنان جدید: {counts['players_new']} | به‌روزشده: {counts['players_updated']}",
        f"🏆 تورنمنت‌های جدید: {counts['tournaments_new']} | به‌روزشده: {counts['tournaments_updated']}",
        f"♟️ مسابقات وارد‌شده: {counts['matches_new']} | ردشده: {counts['matches_skipped']}",
    ]
    if counts["errors"]:
        lines.append(f"\n⚠️ خطاها ({len(counts['errors'])}):")
        for e in counts["errors"][:10]:
            lines.append(f"  • {e}")
        if len(counts["errors"]) > 10:
            lines.append(f"  … و {len(counts['errors']) - 10} خطای دیگر")
    return "\n".join(lines)
